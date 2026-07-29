"""Genera Word (.docx) editable de les fitxes de repaso — disseny tipus ficha del cole."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Union

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from generate_fichas_repaso import (
    AREAS,
    ExerciseBlock,
    FICHAS,
    IMAGES,
    build_mezcla_pages,
    collect_fitxa_files,
    extract_title,
    parse_fitxa_blocks,
)

OUT = Path(__file__).resolve().parent / "salida" / "fichas_repaso"

ContentParent = Union[Document, _Cell]

# Colors (hex without #)
C_BLUE = "1E4DB7"
C_BLUE_LIGHT = "EEF3FF"
C_BORDER = "D5DDEC"
C_CARD_BG = "FAFBFE"
C_TEXT = "1F2530"
C_MUTED = "64748B"
C_TIP_BG = "FFF8E6"
C_TIP_BORDER = "F0D78C"
C_LINE = "94A3B8"

AREA_RGB = {
    "MAT": RGBColor(0x1E, 0x4D, 0xB7),
    "CAS": RGBColor(0xC4, 0x5A, 0x11),
    "CAT": RGBColor(0x2A, 0x7A, 0x3B),
    "MED": RGBColor(0x6B, 0x4C, 0x9A),
    "ANG": RGBColor(0x0D, 0x7A, 0x8C),
}

AREA_HEX = {
    "MAT": "1E4DB7",
    "CAS": "C45A11",
    "CAT": "2A7A3B",
    "MED": "6B4C9A",
    "ANG": "0D7A8C",
}

FONT = "Calibri"
FONT_BODY = 14
FONT_SMALL = 12
FONT_TITLE = 16
FONT_CARD_TITLE = 15

# Cel·les per escriure a mà
ANSWER_BOX_H = Mm(18)
ANSWER_BOX_W = Mm(20)
ANSWER_GRID_H = Mm(16)
ANSWER_GRID_W = Mm(17)

BLANK_RE = re.compile(r"_{3,}")


# ── OXML helpers ──────────────────────────────────────────────────────────────


def set_cell_shading(cell: _Cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_margins(
    cell: _Cell, top: int = 100, start: int = 160, bottom: int = 100, end: int = 160
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)


def set_table_width_pct(table, pct: int = 100) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(pct * 50))
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)


def set_table_borders(
    table, color_hex: str = C_BORDER, sz: int = 6, edges: tuple[str, ...] | None = None
) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    all_edges = edges or ("top", "left", "bottom", "right", "insideH", "insideV")
    for edge in all_edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color_hex)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def set_paragraph_bottom_border(paragraph: Paragraph, color: str = C_LINE, sz: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:color"), color)
    bottom.set(qn("w:space"), "1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_row_min_height(row, height_mm: float) -> None:
    row.height = Mm(height_mm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def add_field_cell(cell: _Cell, label: str) -> None:
    """Etiqueta + línia d'escriptura dins una cel·la."""
    set_cell_margins(cell, 40, 0, 40, 60)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label} ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(FONT_SMALL)
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    p.add_run("\n")
    line_p = cell.add_paragraph()
    line_p.add_run(" ")
    set_paragraph_bottom_border(line_p)
    line_p.paragraph_format.space_after = Pt(0)


def style_paragraph(
    paragraph: Paragraph,
    size: float = FONT_BODY,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    align=None,
    space_before: float = 0,
    space_after: float = 4,
    line_spacing: float = 1.25,
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if align is not None:
        pf.alignment = align
    for run in paragraph.runs:
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color


def new_paragraph(parent: ContentParent, align=None) -> Paragraph:
    p = parent.add_paragraph()
    if align is not None:
        p.paragraph_format.alignment = align
    return p


# ── Fields & lines ────────────────────────────────────────────────────────────


def add_label_line(parent: ContentParent, label: str, width_mm: float = 55) -> None:
    """Nom: ______ amb línia real (taula invisible 2 columnes)."""
    table = parent.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width_pct(table, 55)
    set_table_borders(table, "FFFFFF", sz=0, edges=())
    lbl_cell, line_cell = table.rows[0].cells
    set_cell_margins(lbl_cell, 0, 0, 0, 40)
    set_cell_margins(line_cell, 0, 0, 0, 0)

    p = lbl_cell.paragraphs[0]
    r = p.add_run(f"{label} ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(FONT_SMALL)
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    lp = line_cell.paragraphs[0]
    lp.add_run(" ")
    set_paragraph_bottom_border(lp)
    lp.paragraph_format.space_after = Pt(0)
    line_cell.width = Mm(width_mm)


def add_handwriting_box(parent: ContentParent, unit: str = "") -> None:
    """Quadre gran per escriure la resposta a mà (fora de taules)."""
    box = parent.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width_pct(box, 42)
    set_table_borders(box, C_BORDER, sz=6)
    cell = box.rows[0].cells[0]
    set_cell_shading(cell, "FFFFFF")
    set_cell_margins(cell, 120, 120, 120, 120)
    cell.width = ANSWER_BOX_W
    set_row_min_height(box.rows[0], 18)

    p = cell.paragraphs[0]
    p.add_run(" ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if unit:
        up = new_paragraph(parent)
        up.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ur = up.add_run(unit)
        ur.font.name = FONT
        ur.font.size = Pt(FONT_SMALL)
        ur.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    gap = new_paragraph(parent)
    gap.paragraph_format.space_after = Pt(6)


def fill_answer_cell(cell: _Cell, unit: str = "", grid: bool = False) -> None:
    """Cel·la buida gran dins una taula."""
    set_cell_shading(cell, "FFFFFF")
    set_cell_margins(cell, 100, 90, 100, 90)
    if grid:
        cell.width = ANSWER_GRID_W
    p = cell.paragraphs[0]
    p.add_run(" ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if unit:
        pu = cell.add_paragraph()
        pu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ur = pu.add_run(unit)
        ur.font.name = FONT
        ur.font.size = Pt(FONT_SMALL)
        ur.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def add_answer_line(parent: ContentParent, prefix: str = "") -> None:
    if prefix:
        p = new_paragraph(parent)
        add_formatted_text(p, prefix, size=FONT_BODY)
    add_handwriting_box(parent)


# ── Header & cover ──────────────────────────────────────────────────────────────


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(11)
    section.bottom_margin = Mm(11)
    section.left_margin = Mm(13)
    section.right_margin = Mm(13)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(FONT_BODY)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.2
    return doc


def add_cover(doc: Document) -> None:
    doc.add_paragraph()
    aray = IMAGES / "aray.png"
    if aray.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(aray), width=Mm(32))

    t = doc.add_paragraph("Fichas de repaso")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(t, size=26, bold=True, color=RGBColor(0x1E, 0x4D, 0xB7), space_after=6)

    s = doc.add_paragraph("Aray · 3r primària · Repàs d'estiu")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(s, size=13, color=RGBColor(0x64, 0x74, 0x8B), space_after=4)

    s2 = doc.add_paragraph("Roblox · Arena de Cuchillos")
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(s2, size=11, italic=True, color=RGBColor(0x94, 0xA3, 0xB8), space_after=16)

    note = doc.add_paragraph(
        "Document editable: pots moure exercicis, canviar salts de pàgina i ajustar el disseny."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(note, size=9.5, color=RGBColor(0x94, 0xA3, 0xB8))
    doc.add_page_break()


def add_page_header(doc: Document, page_num: int) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_pct(table, 100)
    set_table_borders(table, C_BLUE, sz=8, edges=("top", "left", "bottom", "right"))
    row = table.rows[0]
    for cell in row.cells:
        set_cell_shading(cell, C_BLUE_LIGHT)
        set_cell_margins(cell, 80, 120, 80, 120)

    c0, c1, c2 = row.cells
    c0.width = Mm(24)
    c2.width = Mm(36)

    aray = IMAGES / "aray.png"
    if aray.exists():
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.add_run().add_picture(str(aray), width=Mm(17))

    # Nom i Data alineats
    inner = c1.add_table(rows=1, cols=2)
    set_table_width_pct(inner, 100)
    set_table_borders(inner, "FFFFFF", sz=0, edges=())
    add_field_cell(inner.rows[0].cells[0], "Nom:")
    add_field_cell(inner.rows[0].cells[1], "Data:")

    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p2.add_run("Repàs mesclat\n")
    r1.font.name = FONT
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r2 = p2.add_run(f"Pàg. {page_num}")
    r2.bold = True
    r2.font.name = FONT
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x1E, 0x4D, 0xB7)

    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(5)


# ── Exercise cards ────────────────────────────────────────────────────────────


def add_exercise_card(doc: Document, ex: ExerciseBlock) -> None:
    outer = doc.add_table(rows=1, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_pct(outer, 100)
    set_table_borders(outer, C_BORDER, sz=8)
    card = outer.rows[0].cells[0]
    set_cell_shading(card, C_CARD_BG)
    set_cell_margins(card, 110, 170, 110, 170)

    # Capçalera de la targeta: badge + àrea
    head = card.add_table(rows=1, cols=2)
    set_table_width_pct(head, 100)
    set_table_borders(head, "FFFFFF", sz=0, edges=())
    badge_cell, meta_cell = head.rows[0].cells
    badge_cell.width = Mm(14)
    set_cell_margins(badge_cell, 0, 0, 0, 80)
    set_cell_margins(meta_cell, 0, 0, 0, 0)

    hex_color = AREA_HEX.get(ex.area_short, C_BLUE)
    set_cell_shading(badge_cell, hex_color)
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(ex.area_short)
    br.bold = True
    br.font.name = FONT
    br.font.size = Pt(9)
    br.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    mp = meta_cell.paragraphs[0]
    mr = mp.add_run(f"{ex.area_label}  ·  Fitxa {ex.fitxa_num}")
    mr.font.name = FONT
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    title = new_paragraph(card)
    tr = title.add_run(ex.section_title)
    tr.bold = True
    tr.font.name = FONT
    tr.font.size = Pt(FONT_CARD_TITLE)
    tr.font.color.rgb = RGBColor(0x1F, 0x25, 0x30)
    title.paragraph_format.space_after = Pt(6)

    add_markdown_body(card, ex.md_body)

    spacer = new_paragraph(doc)
    spacer.paragraph_format.space_after = Pt(7)


# ── Markdown → Word ───────────────────────────────────────────────────────────


def add_image(parent: ContentParent, src: str, width_hint: int | None = None) -> None:
    name = Path(src.replace("\\", "/")).name
    path = IMAGES / name
    if not path.exists():
        return
    max_mm = 82 if (width_hint and width_hint > 200) else 48
    p = new_paragraph(parent, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(str(path), width=Mm(max_mm))


def parse_table(lines: list[str]) -> tuple[list[str] | None, list[list[str]], bool]:
    rows: list[list[str]] = []
    for line in lines:
        if "|" not in line:
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return None, [], False

    header = rows[0]
    is_grid = all(not c or c == ":" for c in header) or all(
        re.match(r"^:?-+:?$", c.replace(" ", "")) for c in header
    )
    if is_grid:
        return None, rows, True
    return header, rows[1:], False


def detect_answer_columns(header: list[str]) -> set[int]:
    cols: set[int] = set()
    for j, h in enumerate(header):
        hl = h.lower().strip("*_ ")
        if not hl:
            continue
        if any(
            k in hl
            for k in (
                "resposta",
                "respuesta",
                "operació",
                "operacion",
                "tipo",
                "sílaba",
                "silaba",
            )
        ):
            cols.add(j)
    return cols


def split_op_answer(text: str) -> tuple[str, str] | None:
    m = BLANK_RE.search(text)
    if not m:
        return None
    before = text[: m.start()].rstrip()
    after = text[m.end() :].strip()
    op = before
    unit = ""
    if before and after:
        op = f"{before}\n{after}"
    elif after and not before:
        unit = after
        op = ""
    return op, unit


def expand_row(
    cells: list[str], answer_cols: set[int] | None, is_grid: bool
) -> list[tuple[str, str, str]]:
    """(tipus, text, unit) — tipus: op | ans | txt | hdr"""
    out: list[tuple[str, str, str]] = []
    for j, raw in enumerate(cells):
        text = raw.strip()
        if not text:
            if answer_cols and j in answer_cols:
                out.append(("ans", "", ""))
            continue

        if BLANK_RE.fullmatch(text):
            out.append(("ans", "", ""))
            continue

        split = split_op_answer(text)
        if split:
            op, unit = split
            if is_grid:
                if op:
                    out.append(("op", op, ""))
                out.append(("ans", "", unit))
            else:
                if op:
                    out.append(("txt", op, ""))
                out.append(("ans", "", unit))
            continue

        if answer_cols and j in answer_cols:
            out.append(("ans", "", ""))
            continue

        out.append(("txt", text, ""))
    return out


def add_md_table(
    parent: ContentParent, header: list[str] | None, body: list[list[str]], is_grid: bool
) -> None:
    answer_cols = detect_answer_columns(header) if header else set()

    expanded_rows: list[list[tuple[str, str, str]]] = []
    if header:
        expanded_rows.append([("hdr", h, "") for h in header])
    for row in body:
        expanded_rows.append(expand_row(row, answer_cols if header else None, is_grid))

    # En graella de mates, expandir operació + resposta per cel·la
    if is_grid:
        expanded_rows = []
        for row in body:
            new_row: list[tuple[str, str, str]] = []
            for cell in row:
                if not cell.strip():
                    continue
                split = split_op_answer(cell.strip())
                if split:
                    op, unit = split
                    if op:
                        new_row.append(("op", op, ""))
                    new_row.append(("ans", "", unit))
                else:
                    new_row.append(("txt", cell.strip(), ""))
            if new_row:
                expanded_rows.append(new_row)

    if not expanded_rows:
        return

    cols = max(len(r) for r in expanded_rows)
    table = parent.add_table(rows=len(expanded_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_pct(table, 100)
    set_table_borders(table, C_BORDER, sz=5)

    for i, row_data in enumerate(expanded_rows):
        row = table.rows[i]
        has_answer = any(c[0] == "ans" for c in row_data)
        if has_answer:
            set_row_min_height(row, 18 if is_grid else 20)

        for j in range(cols):
            cell = row.cells[j]
            if j >= len(row_data):
                fill_answer_cell(cell, grid=is_grid)
                continue

            kind, text, unit = row_data[j]

            if kind == "hdr":
                set_cell_shading(cell, C_BLUE_LIGHT)
                set_cell_margins(cell, 70, 90, 70, 90)
                cp = cell.paragraphs[0]
                add_formatted_text(cp, text, size=FONT_SMALL, bold=True)
                continue

            if kind == "ans":
                fill_answer_cell(cell, unit, grid=is_grid)
                continue

            if kind == "op":
                set_cell_margins(cell, 80, 90, 80, 90)
                set_cell_shading(cell, C_BLUE_LIGHT)
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_formatted_text(cp, text, size=FONT_BODY, bold=False)
                continue

            set_cell_margins(cell, 80, 90, 80, 90)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_formatted_text(cp, text, size=FONT_BODY)

    gap = new_paragraph(parent)
    gap.paragraph_format.space_after = Pt(6)


def add_tip_box(parent: ContentParent, text: str) -> None:
    box = parent.add_table(rows=1, cols=1)
    set_table_width_pct(box, 100)
    set_table_borders(box, C_TIP_BORDER, sz=6)
    cell = box.rows[0].cells[0]
    set_cell_shading(cell, C_TIP_BG)
    set_cell_margins(cell, 70, 120, 70, 120)
    p = cell.paragraphs[0]
    r = p.add_run("Recorda: ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(FONT_SMALL)
    add_formatted_text(p, text.replace("**Recorda:**", "").strip(), size=FONT_SMALL, italic=True)


def add_markdown_body(parent: ContentParent, md: str) -> None:
    lines = md.replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            if not in_code and code_lines:
                box = parent.add_table(rows=1, cols=1)
                set_table_width_pct(box, 100)
                set_table_borders(box, C_BORDER, sz=4)
                cell = box.rows[0].cells[0]
                set_cell_shading(cell, "F8FAFC")
                set_cell_margins(cell, 80, 120, 80, 120)
                p = cell.paragraphs[0]
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(FONT_SMALL)
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("**Recorda:**"):
            add_tip_box(parent, stripped)
            i += 1
            continue

        img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)(?:\s+width=\"(\d+)\")?", stripped)
        if img:
            hint = int(img.group(3)) if img.group(3) else None
            if not hint:
                w = re.search(r'width="(\d+)"', stripped)
                hint = int(w.group(1)) if w else None
            add_image(parent, img.group(2), hint)
            i += 1
            continue

        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            box = parent.add_table(rows=1, cols=1)
            set_table_width_pct(box, 100)
            set_table_borders(box, "E2E8F0", sz=4)
            cell = box.rows[0].cells[0]
            set_cell_shading(cell, "F8FAFC")
            set_cell_margins(cell, 80, 140, 80, 140)
            p = cell.paragraphs[0]
            add_formatted_text(p, text, size=FONT_BODY, italic=True)
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and "|" in lines[i]:
                block.append(lines[i])
                i += 1
            hdr, body, is_grid = parse_table(block)
            add_md_table(parent, hdr, body, is_grid)
            continue

        if re.match(r"^_{8,}$", stripped):
            add_answer_line(parent)
            i += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m_num:
            rest = m_num.group(2)
            p = new_paragraph(parent)
            p.paragraph_format.left_indent = Mm(2)
            p.paragraph_format.first_line_indent = Mm(-2)
            num_run = p.add_run(f"{m_num.group(1)}. ")
            num_run.bold = True
            num_run.font.name = FONT
            num_run.font.size = Pt(FONT_BODY)
            if BLANK_RE.search(rest):
                before = BLANK_RE.sub("", rest, count=1).strip()
                m_unit = re.search(r"_{3,}\s*(.*)$", rest)
                unit = m_unit.group(1).strip() if m_unit else ""
                if before:
                    add_formatted_text(p, before, size=FONT_BODY)
                add_handwriting_box(parent, unit)
            else:
                add_formatted_text(p, rest, size=FONT_BODY)
            i += 1
            if i < len(lines) and re.match(r"^_{6,}$", lines[i].strip()):
                add_handwriting_box(parent)
                i += 1
            continue

        m_check = re.match(r"^[-*]\s+\[([ x])\]\s+(.*)", stripped)
        if m_check:
            p = new_paragraph(parent)
            p.paragraph_format.left_indent = Mm(3)
            sym = "☑" if m_check.group(1).lower() == "x" else "☐"
            add_formatted_text(p, f"{sym}  {m_check.group(2)}", size=FONT_BODY)
            i += 1
            continue

        m_bullet = re.match(r"^[-*]\s+(.*)", stripped)
        if m_bullet:
            p = new_paragraph(parent)
            p.paragraph_format.left_indent = Mm(4)
            p.style = "List Bullet"
            add_formatted_text(p, m_bullet.group(1), size=FONT_BODY)
            i += 1
            continue

        if stripped.startswith("### "):
            p = new_paragraph(parent)
            add_formatted_text(p, stripped[4:], size=FONT_BODY, bold=True)
            i += 1
            continue

        if BLANK_RE.search(stripped):
            before = BLANK_RE.sub("", stripped, count=1).strip()
            m_unit = re.search(r"_{3,}\s*(.*)$", stripped)
            unit = m_unit.group(1).strip() if m_unit else ""
            if before:
                p = new_paragraph(parent)
                add_formatted_text(p, before, size=FONT_BODY)
            add_handwriting_box(parent, unit)
            if i + 1 < len(lines) and re.match(r"^_{6,}$", lines[i + 1].strip()):
                i += 2
                continue
            i += 1
            continue

        p = new_paragraph(parent)
        add_formatted_text(p, stripped, size=FONT_BODY)
        if i + 1 < len(lines) and re.match(r"^_{6,}$", lines[i + 1].strip()):
            i += 2
            add_handwriting_box(parent)
            continue
        i += 1


def add_formatted_text(
    paragraph: Paragraph,
    text: str,
    size: float = FONT_BODY,
    bold: bool = False,
    italic: bool = False,
) -> None:
    clean = BLANK_RE.sub("", text).strip()
    if not clean:
        return
    parts = re.split(r"(\*\*[^*]+\*\*)", clean)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.italic = italic
        if bold and not (part.startswith("**") and part.endswith("**")):
            run.bold = True


# ── Documents ─────────────────────────────────────────────────────────────────


def build_mezcla_document() -> Document:
    doc = setup_document()
    add_cover(doc)

    pages = build_mezcla_pages()
    for num, page_blocks in enumerate(pages, start=1):
        add_page_header(doc, num)
        for ex in page_blocks:
            add_exercise_card(doc, ex)
        if num < len(pages):
            doc.add_page_break()
    return doc


def build_asignaturas_document() -> Document:
    doc = setup_document()
    add_cover(doc)

    folder_to_short = {
        "01_mates": "MAT",
        "02_castellano": "CAS",
        "03_catala": "CAT",
        "04_medi": "MED",
        "05_angles": "ANG",
    }

    for area_folder, area_label, _ in AREAS:
        short = folder_to_short[area_folder]
        h = doc.add_heading(area_label, level=1)
        if h.runs:
            h.runs[0].font.color.rgb = AREA_RGB.get(short, RGBColor(0, 0, 0))
            h.runs[0].font.name = FONT

        for md_path in collect_fitxa_files(FICHAS / area_folder):
            title = extract_title(md_path.read_text(encoding="utf-8"))
            ht = doc.add_heading(title, level=2)
            if ht.runs:
                ht.runs[0].font.name = FONT
                ht.runs[0].font.size = Pt(14)

            for ex in parse_fitxa_blocks(md_path, area_folder, area_label):
                add_exercise_card(doc, ex)

        doc.add_page_break()
    return doc


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)

    mezcla = build_mezcla_document()
    mezcla_path = OUT / "fichas_mezcla.docx"
    try:
        mezcla.save(mezcla_path)
        print(f"DOCX -> {mezcla_path}")
    except PermissionError:
        alt = OUT / "fichas_mezcla_nou.docx"
        mezcla.save(alt)
        print(f"DOCX -> {alt}  (tanca l'antic a Word i torna a generar)")

    por_area = build_asignaturas_document()
    area_path = OUT / "fichas_por_asignatura.docx"
    try:
        por_area.save(area_path)
        print(f"DOCX -> {area_path}")
    except PermissionError:
        alt = OUT / "fichas_por_asignatura_nou.docx"
        por_area.save(alt)
        print(f"DOCX -> {alt}  (tanca l'antic a Word i torna a generar)")


if __name__ == "__main__":
    main()
